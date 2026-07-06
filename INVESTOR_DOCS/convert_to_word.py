"""
CADRO HR — Yatırımcı Belgelerini Word'e Dönüştür
Markdown → .docx (python-docx)
Çalıştır: python3 INVESTOR_DOCS/convert_to_word.py
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "word_output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

DOCS = [
    "01_Problem_Statement.md",
    "02_Product_Features.md",
    "03_Competitor_Analysis.md",
    "04_Product_Roadmap.md",
    "05_Technical_Architecture.md",
    "06_Pricing_Model.md",
    "07_Fundraising_Strategy.md",
    "08_Data_Room_Checklist.md",
    "09_Investor_Outreach_Playbook.md",
    "10_Due_Diligence_QA.md",
    "11_Closing_Execution_Plan.md",
    "12_Master_Investor_Pack.md",
    "13_Target_Investor_List.md",
    "14_Pitch_Deck_Outline.md",
    "15_Financial_Model_Assumptions.md",
    "16_Investor_One_Pager.md",
    "17_Meeting_Script_Demo_Flow.md",
    "18_Sale_Pitch_Deck.md",
    "19_CEO_Presentation_Script.md",
    "20_QA_Meeting_Version.md",
    "21_Investor_Structure_Plan.md",
    "05_Technical_Architecture_EN_Redacted.md",
]

# Renkler
COLOR_BRAND   = RGBColor(0x0D, 0x9E, 0xDF)   # CADRO mavi
COLOR_H1      = RGBColor(0x1E, 0x3A, 0x5F)   # koyu lacivert
COLOR_H2      = RGBColor(0x0D, 0x9E, 0xDF)   # mavi
COLOR_H3      = RGBColor(0x33, 0x33, 0x33)   # koyu gri
COLOR_BLACK   = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_TABLE_H = RGBColor(0x0D, 0x9E, 0xDF)   # tablo başlık arka plan
COLOR_TABLE_A = RGBColor(0xF0, 0xF7, 0xFF)   # tablo alt satır

def set_cell_bg(cell, hex_color):
    """Tablo hücresi arka plan rengi."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    """Yatay çizgi ekle."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0D9EDF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_cover_page(doc, filename):
    """Kapak sayfası."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Logo yerine şirket adı
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CADRO HR")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = COLOR_BRAND

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Yatırımcı Dokümanı")
    r2.font.size = Pt(14)
    r2.font.color.rgb = COLOR_H1

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Belge adı
    title_map = {
        "01_Problem_Statement.md":    "Problem Tanımı & Pazar Fırsatı",
        "02_Product_Features.md":     "Ürün Özeti & Özellik Kataloğu",
        "03_Competitor_Analysis.md":  "Rakip Analizi",
        "04_Product_Roadmap.md":      "Ürün Yol Haritası",
        "05_Technical_Architecture.md": "Teknik Mimari & Platform Analizi",
        "06_Pricing_Model.md":        "Fiyatlandırma Modeli & Gelir Projeksiyonu",
    }
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(title_map.get(filename, filename))
    r3.bold = True
    r3.font.size = Pt(22)
    r3.font.color.rgb = COLOR_H1

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Gizli — Yalnızca Yatırımcı Kullanımı  |  Mayıs 2026")
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def set_paragraph_keep_together(paragraph, keep_with_next=False):
    """Başlığın sayfa kırılmasıyla bölünmesini engelle."""
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_with_next

def parse_inline(run_parent, text):
    """Bold (**text**) ve code (`text`) inline biçimlendirme."""
    # Önce bold, sonra code
    pattern = re.compile(r'\*\*(.*?)\*\*|`([^`]+)`')
    last = 0
    for m in pattern.finditer(text):
        # öncesi düz metin
        if m.start() > last:
            r = run_parent.add_run(text[last:m.start()])
            r.font.color.rgb = COLOR_BLACK
        if m.group(1) is not None:
            r = run_parent.add_run(m.group(1))
            r.bold = True
            r.font.color.rgb = COLOR_BLACK
        else:
            r = run_parent.add_run(m.group(2))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        last = m.end()
    if last < len(text):
        r = run_parent.add_run(text[last:])
        r.font.color.rgb = COLOR_BLACK

def add_table_from_md(doc, header_row, rows):
    """Markdown tablosunu Word tablosuna dönüştür."""
    col_count = len(header_row)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    # Başlık satırı
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = h.strip().strip('*')
        set_cell_bg(hdr_cells[i], '0D9EDF')
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    # Veri satırları
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        bg = 'F0F7FF' if ri % 2 == 0 else 'FFFFFF'
        for i, cell_val in enumerate(row[:col_count]):
            val = cell_val.strip()
            cells[i].text = val
            set_cell_bg(cells[i], bg)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = COLOR_BLACK
    doc.add_paragraph()

def convert_md_to_docx(md_path, output_path):
    doc = Document()
    # Kenar boşlukları
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    filename = os.path.basename(md_path)
    add_cover_page(doc, filename)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    table_header = None
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip('\n')

        # Kod bloğu başlangıç/bitiş
        if raw.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                p = doc.add_paragraph()
                for cl in code_lines:
                    r = p.add_run(cl + '\n')
                    r.font.name = 'Courier New'
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                i += 1
                continue
        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # Tablo
        if raw.strip().startswith('|'):
            parts = raw.split('|')
            # Baştaki/sondaki boş elemanları kaldır, ortadaki boşlukları koru
            if parts and parts[0].strip() == '':
                parts = parts[1:]
            if parts and parts[-1].strip() == '':
                parts = parts[:-1]
            cols = parts if parts else []
            # Ayırıcı satır mı?
            if all(re.match(r'^[\s:\-]+$', c) for c in cols):
                i += 1
                continue
            if table_header is None:
                table_header = cols
            else:
                table_rows.append(cols)
            # Sonraki satır tablo mu?
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # Tabloyu yaz
                add_table_from_md(doc, table_header, table_rows)
                table_header = None
                table_rows = []
                i += 1
                continue

        # Boş satır
        if not raw.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Başlıklar
        h_match = re.match(r'^(#{1,4})\s+(.*)', raw)
        if h_match:
            level = len(h_match.group(1))
            text  = h_match.group(2).strip()
            p = doc.add_paragraph()
            set_paragraph_keep_together(p, keep_with_next=True)
            r = p.add_run(text)
            r.bold = True
            if level == 1:
                r.font.size = Pt(20)
                r.font.color.rgb = COLOR_H1
                p.paragraph_format.space_before = Pt(18)
                add_horizontal_rule(doc)
            elif level == 2:
                r.font.size = Pt(15)
                r.font.color.rgb = COLOR_H2
                p.paragraph_format.space_before = Pt(12)
            elif level == 3:
                r.font.size = Pt(12)
                r.font.color.rgb = COLOR_H3
                p.paragraph_format.space_before = Pt(8)
            else:
                r.font.size = Pt(11)
                r.font.color.rgb = COLOR_H3
                p.paragraph_format.space_before = Pt(6)
            i += 1
            continue

        # Yatay çizgi
        if re.match(r'^---+\s*$', raw):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Madde işareti
        bullet_match = re.match(r'^(\s*)([-*•]|\d+\.)\s+(.*)', raw)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text   = bullet_match.group(3)
            style  = 'List Bullet' if not bullet_match.group(2)[0].isdigit() else 'List Number'
            p = doc.add_paragraph(style=style)
            if bullet_match.group(2)[0].isdigit() and indent == 0:
                set_paragraph_keep_together(p, keep_with_next=True)
            p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
            parse_inline(p, text)
            i += 1
            continue

        # Alıntı (blockquote)
        if raw.strip().startswith('>'):
            text = raw.strip().lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.italic = True
            i += 1
            continue

        # Kalın başlık satırları (**...**)
        if raw.strip().startswith('**') and raw.strip().endswith('**'):
            p = doc.add_paragraph()
            r = p.add_run(raw.strip().strip('*'))
            r.bold = True
            r.font.color.rgb = COLOR_H3
            i += 1
            continue

        # Düz paragraf
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        parse_inline(p, raw.strip())
        i += 1

    doc.save(output_path)
    print(f"  ✅ {os.path.basename(output_path)}")


if __name__ == "__main__":
    print("\n🔄 CADRO — Markdown → Word Dönüşümü Başlıyor...\n")
    for doc_file in DOCS:
        md_path = os.path.join(BASE_DIR, doc_file)
        if not os.path.exists(md_path):
            print(f"  ⚠️  Atlandı (bulunamadı): {doc_file}")
            continue
        out_name = doc_file.replace('.md', '.docx')
        out_path = os.path.join(OUTPUT_DIR, out_name)
        try:
            convert_md_to_docx(md_path, out_path)
        except Exception as e:
            print(f"  ❌ {doc_file} — HATA: {e}")
    print(f"\n✅ Tüm belgeler '{OUTPUT_DIR}' klasörüne kaydedildi.\n")
